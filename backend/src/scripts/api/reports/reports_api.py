#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
广东省空气质量监测系统 - 报告生成 API 服务
提供报告生成、下载和历史记录查询功能
"""

import os
import sys
import json
import time
import logging
import traceback
import io
from datetime import datetime, timedelta
from pathlib import Path
from flask import Flask, jsonify, request, send_file, abort, Response
from flask_cors import CORS
import pandas as pd
import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib.font_manager import FontProperties
import seaborn as sns
import mysql.connector
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.units import inch, mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import mimetypes
import uuid
from urllib.parse import quote
import requests
from dotenv import load_dotenv

# Load environment variables from .env file
backend_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))))
dotenv_path = os.path.join(backend_dir, '.env')
if os.path.exists(dotenv_path):
    load_dotenv(dotenv_path)

# 获取当前脚本目录
current_dir = os.path.dirname(os.path.abspath(__file__))
# 获取项目根目录
project_root = str(Path(__file__).resolve().parents[4])
# 添加项目根目录到Python路径
sys.path.append(project_root)

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(message)s',
    handlers=[
        logging.StreamHandler(sys.stdout),
        logging.FileHandler(os.path.join(project_root, 'src', 'scripts', 'api', 'logs', 'reports_api.log'), encoding='utf-8')
    ]
)
logger = logging.getLogger('reports_api')

# 导入数据处理和绘图工具
try:
    from src.scripts.utils.data_utils import load_data, process_data
    from src.scripts.utils.plot_utils import create_bar_chart, create_line_chart
    # 导入报告生成模块
    from . import report_generation
except ImportError as e:
    logger.warning(f"导入工具模块失败: {str(e)}")
    # 创建简单的替代函数，以防导入失败
    def load_data(*args, **kwargs):
        return pd.DataFrame()
    
    def process_data(*args, **kwargs):
        return pd.DataFrame()
    
    def create_bar_chart(*args, **kwargs):
        return None
    
    def create_line_chart(*args, **kwargs):
        return None
    
    # 创建空的report_generation模块，以防导入失败
    class ReportGenerationMock:
        @staticmethod
        def generate_report(*args, **kwargs):
            return {'success': False, 'error': '报告生成模块导入失败'}
    
    report_generation = ReportGenerationMock()

# 创建Flask应用
app = Flask(__name__)
CORS(app)

# 报告存储目录
REPORTS_DIR = os.path.join(project_root, 'data', 'reports')
os.makedirs(REPORTS_DIR, exist_ok=True)

# 报告子目录（按类型）
REPORTS_PDF_DIR = os.path.join(REPORTS_DIR, 'pdf')
REPORTS_EXCEL_DIR = os.path.join(REPORTS_DIR, 'excel')
REPORTS_WORD_DIR = os.path.join(REPORTS_DIR, 'word')
REPORTS_HTML_DIR = os.path.join(REPORTS_DIR, 'html')

# 创建所有报告目录
for dir_path in [REPORTS_DIR, REPORTS_PDF_DIR, REPORTS_EXCEL_DIR, REPORTS_WORD_DIR, REPORTS_HTML_DIR]:
    os.makedirs(dir_path, exist_ok=True)
    logger.info(f"确保报告目录存在: {dir_path}")

# 字体资源目录
FONTS_DIR = os.path.join(project_root, 'data', 'fonts')
os.makedirs(FONTS_DIR, exist_ok=True)

# 默认字体文件路径
DEFAULT_FONT_PATH = os.path.join(FONTS_DIR, 'simhei.ttf')

# 报告元数据存储路径
REPORT_METADATA_PATH = os.path.join(REPORTS_DIR, 'reports_metadata.json')

# 数据库配置
DB_CONFIG = {
    'host': os.environ.get('DB_HOST', 'localhost'),
    'user': os.environ.get('DB_USER', 'root'),
    'password': os.environ.get('DB_PASSWORD'),
    'database': os.environ.get('DB_NAME', 'air_quality_monitoring'),
    'port': int(os.environ.get('DB_PORT', '3306')),
    'connect_timeout': 10,
    'connection_timeout': 10
}

# Validate that password is loaded
if not DB_CONFIG['password']:
    raise ValueError("DB_PASSWORD environment variable is required but not set")

# 简化的错误处理函数
def handle_error(error_msg, status_code=500):
    """处理API错误并返回标准格式的错误响应
    
    Args:
        error_msg: 错误信息字符串或异常对象
        status_code: HTTP状态码，默认为500
        
    Returns:
        Flask响应对象，包含错误信息和状态码
    """
    # 记录错误
    if isinstance(error_msg, Exception):
        logger.error(f"错误: {str(error_msg)}\n{traceback.format_exc()}")
        error_msg = str(error_msg)
    else:
        logger.error(error_msg)
    
    # 构建响应
    try:
        response = {
            'error': error_msg, 
            'status': 'error', 
            'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        }
        return jsonify(response), status_code
    except Exception as e:
        # 如果JSON序列化失败，返回纯文本错误
        logger.error(f"构建错误响应失败: {str(e)}")
        return f"服务器错误: {error_msg}", status_code

# 加载空气质量数据
def load_air_quality_data(start_date, end_date, region):
    conn = None
    cursor = None
    try:
        today = datetime.now().strftime('%Y-%m-%d')
        
        # 记录请求参数，用于排查问题
        logger.info(f"加载空气质量数据 - 参数: start_date={start_date}, end_date={end_date}, region={region}")
        
        # 准备返回的数据框架
        aqi_data = pd.DataFrame()
        pollutants = pd.DataFrame()
        
        # 首先尝试从数据库加载所有数据，不管是否包括今天
        # 连接数据库
        conn = mysql.connector.connect(**DB_CONFIG)
        cursor = conn.cursor(dictionary=True)
        
        # 构建查询条件
        region_condition = ""
        region_params = []
        
        if region != 'all':
            # 更宽松的区域查询逻辑 - 模糊匹配或精确匹配
            region_condition = "AND (city = %s OR city LIKE %s OR city LIKE %s)"
            region_params = [region, f"{region}%", f"%{region}%"]
        
        # 查询AQI数据
        aqi_query = f"""
        SELECT record_date as date, aqi_index as aqi, quality_level, city
        FROM air_quality_data
        WHERE record_date BETWEEN %s AND %s {region_condition}
        ORDER BY record_date
        """
        aqi_params = [start_date, end_date] + region_params
        
        logger.info(f"执行AQI查询: {aqi_query}, 参数: {aqi_params}")
        cursor.execute(aqi_query, aqi_params)
        db_aqi_data = cursor.fetchall()
        logger.info(f"AQI查询结果: 获取到 {len(db_aqi_data)} 条记录")
        
        if db_aqi_data:
            aqi_data = pd.DataFrame(db_aqi_data)
            logger.info(f"AQI数据非空: {len(aqi_data)}条, 列: {list(aqi_data.columns)}")
            if not aqi_data.empty:
                logger.info(f"AQI数据样例: {aqi_data.head(1).to_dict('records')}")
        
        # 查询污染物数据
        poll_query = f"""
        SELECT record_date as date, pm25_avg as pm25, pm10_avg as pm10, so2_avg as so2, no2_avg as no2, o3_avg as o3, co_avg as co, city
        FROM air_quality_data
        WHERE record_date BETWEEN %s AND %s {region_condition}
        ORDER BY record_date
        """
        poll_params = [start_date, end_date] + region_params
        
        logger.info(f"执行污染物查询: {poll_query}, 参数: {poll_params}")
        cursor.execute(poll_query, poll_params)
        db_pollutants = cursor.fetchall()
        logger.info(f"污染物查询结果: 获取到 {len(db_pollutants)} 条记录")
        
        if db_pollutants:
            pollutants = pd.DataFrame(db_pollutants)
            logger.info(f"污染物数据非空: {len(pollutants)}条, 列: {list(pollutants.columns)}")
            if not pollutants.empty:
                logger.info(f"污染物数据样例: {pollutants.head(1).to_dict('records')}")
        
        # 检查并尝试获取实时数据
        if start_date <= today <= end_date:
            logger.info(f"时间范围包括今天，尝试获取实时数据")
            try:
                # 获取实时数据
                realtime_data = fetch_realtime_data('all')
                if realtime_data:
                    logger.info(f"成功获取到实时数据: {len(realtime_data)}条")
                    
                    # 转换为DataFrame并与历史数据合并
                    realtime_aqi_records = []
                    realtime_pollutant_records = []
                    
                    for city_data in realtime_data:
                        city_name = city_data.get('name', '')
                        
                        # 如果指定了特定区域但不是全省，使用更宽松的匹配
                        if region != 'all':
                            # 更宽松的匹配
                            if not (region.lower() in city_name.lower() or 
                                   city_name.lower() in region.lower() or
                                   city_name.lower() == region.lower()):
                                continue
                        
                        # 创建AQI记录
                        aqi_record = {
                            'date': today,
                            'aqi': city_data.get('aqi', 0),
                            'quality_level': city_data.get('level', ''),
                            'city': city_name
                        }
                        realtime_aqi_records.append(aqi_record)
                        
                        # 创建污染物记录
                        pollutant_record = {
                            'date': today,
                            'pm25': city_data.get('pm25', 0),
                            'pm10': city_data.get('pm10', 0),
                            'so2': city_data.get('so2', 0),
                            'no2': city_data.get('no2', 0),
                            'o3': city_data.get('o3', 0),
                            'co': float(city_data.get('co', 0)),
                            'city': city_name
                        }
                        realtime_pollutant_records.append(pollutant_record)
                    
                    logger.info(f"处理后的实时数据记录: AQI={len(realtime_aqi_records)}条, 污染物={len(realtime_pollutant_records)}条")
                    
                    # 转换为DataFrame并与历史数据合并
                    if realtime_aqi_records:
                        realtime_aqi_df = pd.DataFrame(realtime_aqi_records)
                        if not aqi_data.empty:
                            aqi_data = pd.concat([aqi_data, realtime_aqi_df], ignore_index=True)
                            logger.info(f"合并后AQI数据: {len(aqi_data)}条")
                        else:
                            aqi_data = realtime_aqi_df
                            logger.info(f"仅使用实时AQI数据: {len(aqi_data)}条")
                    
                    if realtime_pollutant_records:
                        realtime_pollutants_df = pd.DataFrame(realtime_pollutant_records)
                        if not pollutants.empty:
                            pollutants = pd.concat([pollutants, realtime_pollutants_df], ignore_index=True)
                            logger.info(f"合并后污染物数据: {len(pollutants)}条")
                        else:
                            pollutants = realtime_pollutants_df
                            logger.info(f"仅使用实时污染物数据: {len(pollutants)}条")
            except Exception as e:
                logger.error(f"获取实时数据失败: {str(e)}\n{traceback.format_exc()}")
                # 实时数据获取失败，继续使用已经获取的历史数据
        
        # 如果还是没有数据，尝试扩大时间范围再查询
        if aqi_data.empty and pollutants.empty:
            logger.warning(f"在指定时间范围内未找到数据，尝试扩大查询范围")
            
            # 计算一个更长的时间范围 (例如前后延长7天)
            extended_start = (datetime.strptime(start_date, '%Y-%m-%d') - timedelta(days=7)).strftime('%Y-%m-%d')
            extended_end = (datetime.strptime(end_date, '%Y-%m-%d') + timedelta(days=7)).strftime('%Y-%m-%d')
            
            # 重新连接数据库(如果连接已关闭)
            if not conn or not conn.is_connected():
                conn = mysql.connector.connect(**DB_CONFIG)
                cursor = conn.cursor(dictionary=True)
            
            # 重新查询AQI数据
            ext_aqi_query = f"""
            SELECT record_date as date, aqi_index as aqi, quality_level, city
            FROM air_quality_data
            WHERE record_date BETWEEN %s AND %s {region_condition}
            ORDER BY record_date
            """
            ext_aqi_params = [extended_start, extended_end] + region_params
            
            logger.info(f"扩展时间范围AQI查询: {ext_aqi_query}, 参数: {ext_aqi_params}")
            cursor.execute(ext_aqi_query, ext_aqi_params)
            ext_aqi_data = cursor.fetchall()
            logger.info(f"扩展时间范围AQI查询结果: 获取到 {len(ext_aqi_data)} 条记录")
            
            if ext_aqi_data:
                aqi_data = pd.DataFrame(ext_aqi_data)
                logger.info(f"扩展时间范围AQI数据非空: {len(aqi_data)}条, 列: {list(aqi_data.columns)}")
                if not aqi_data.empty:
                    logger.info(f"扩展AQI数据样例: {aqi_data.head(1).to_dict('records')}")
            
            # 重新查询污染物数据
            ext_poll_query = f"""
            SELECT record_date as date, pm25_avg as pm25, pm10_avg as pm10, so2_avg as so2, no2_avg as no2, o3_avg as o3, co_avg as co, city
            FROM air_quality_data
            WHERE record_date BETWEEN %s AND %s {region_condition}
            ORDER BY record_date
            """
            ext_poll_params = [extended_start, extended_end] + region_params
            
            logger.info(f"扩展时间范围污染物查询: {ext_poll_query}, 参数: {ext_poll_params}")
            cursor.execute(ext_poll_query, ext_poll_params)
            ext_pollutants = cursor.fetchall()
            logger.info(f"扩展时间范围污染物查询结果: 获取到 {len(ext_pollutants)} 条记录")
            
            if ext_pollutants:
                pollutants = pd.DataFrame(ext_pollutants)
                logger.info(f"扩展时间范围污染物数据非空: {len(pollutants)}条, 列: {list(pollutants.columns)}")
                if not pollutants.empty:
                    logger.info(f"扩展污染物数据样例: {pollutants.head(1).to_dict('records')}")
        
        # 如果使用了扩展时间范围后有数据，过滤保留原来时间范围内的数据
        if not aqi_data.empty:
            # 检查是否需要过滤 (如果进行了扩展时间范围查询)
            if 'date' in aqi_data.columns:
                # 过滤保留原始时间范围内的数据
                aqi_data = aqi_data[(aqi_data['date'] >= start_date) & (aqi_data['date'] <= end_date)]
                logger.info(f"过滤后AQI数据: {len(aqi_data)}条")
        
        if not pollutants.empty:
            # 同样过滤污染物数据
            if 'date' in pollutants.columns:
                pollutants = pollutants[(pollutants['date'] >= start_date) & (pollutants['date'] <= end_date)]
                logger.info(f"过滤后污染物数据: {len(pollutants)}条")
        
        # 最终返回数据
        # 即使没有查到数据，也返回空的DataFrame而不是错误
        return {
            'aqi_data': aqi_data if not aqi_data.empty else pd.DataFrame(),
            'pollutants': pollutants if not pollutants.empty else pd.DataFrame()
        }
    
    except Exception as e:
        logger.error(f"数据加载失败: {str(e)}\n{traceback.format_exc()}")
        return {'aqi_data': pd.DataFrame(), 'pollutants': pd.DataFrame(), 'error': str(e)}
    finally:
        # 确保资源释放
        if cursor:
            cursor.close()
        if conn and conn.is_connected():
            conn.close()

# 从实时API获取数据
def fetch_realtime_data(region='all'):
    """从实时数据API获取数据
    
    Args:
        region (str): 区域ID，默认为'all'表示全省
        
    Returns:
        list: 包含城市实时数据的列表
    """
    try:
        # 根据区域决定API端点
        if region == 'all':
            endpoint = '/api/province'  # 获取全省数据
        else:
            endpoint = f'/api/realtime/{region}'  # 获取特定城市数据
            
        # 构建请求URL
        url = f"http://localhost:5001{endpoint}"
        
        # 添加参数，禁用模拟数据
        params = {'use_real_data': 'true', 'disable_simulation': 'true'}
            
        # 发送GET请求
        logger.info(f"从实时API获取数据: {url}")
        response = requests.get(url, params=params, timeout=10)
        
        # 检查响应状态
        if response.status_code != 200:
            logger.error(f"实时API返回错误状态码: {response.status_code}")
            return []
            
        # 解析响应数据
        data = response.json()
        
        # 处理不同的响应格式
        if isinstance(data, dict) and 'data' in data:
            result = data['data']
        elif isinstance(data, list):
            result = data
        else:
            result = []
            
        logger.info(f"成功获取实时数据: {len(result)} 条记录")
        return result
    except Exception as e:
        logger.error(f"从实时API获取数据失败: {str(e)}")
        return []

# 加载报告元数据
def load_report_metadata():
    try:
        if os.path.exists(REPORT_METADATA_PATH):
            try:
                with open(REPORT_METADATA_PATH, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    # 验证数据格式
                    if not isinstance(data, list):
                        logger.warning(f"报告元数据格式不正确: {REPORT_METADATA_PATH}，重置为空列表")
                        return []
                    return data
            except json.JSONDecodeError:
                logger.error(f"报告元数据文件格式错误: {REPORT_METADATA_PATH}")
                # 文件存在但格式错误，创建备份
                backup_path = f"{REPORT_METADATA_PATH}.bak.{int(time.time())}"
                try:
                    os.rename(REPORT_METADATA_PATH, backup_path)
                    logger.info(f"已创建元数据文件备份: {backup_path}")
                except Exception as e:
                    logger.error(f"创建备份文件失败: {str(e)}")
                
                # 重新创建文件
                with open(REPORT_METADATA_PATH, 'w', encoding='utf-8') as f:
                    json.dump([], f)
                return []
        else:
            # 如果文件不存在，确保目录存在并创建空文件
            try:
                parent_dir = os.path.dirname(REPORT_METADATA_PATH)
                if not os.path.exists(parent_dir):
                    os.makedirs(parent_dir, exist_ok=True)
                    logger.info(f"创建报告元数据目录: {parent_dir}")
                
                with open(REPORT_METADATA_PATH, 'w', encoding='utf-8') as f:
                    json.dump([], f)
                logger.info(f"创建新的报告元数据文件: {REPORT_METADATA_PATH}")
                return []
            except Exception as e:
                logger.error(f"创建报告元数据文件失败: {str(e)}\n{traceback.format_exc()}")
                return []
    except Exception as e:
        logger.error(f"加载报告元数据失败: {str(e)}\n{traceback.format_exc()}")
        return []

# 保存报告元数据
def save_report_metadata(metadata):
    if metadata is None:
        logger.warning("保存报告元数据: metadata参数为None，将使用空列表")
        metadata = []
        
    if not isinstance(metadata, list):
        logger.warning(f"保存报告元数据: metadata不是列表格式({type(metadata)})，将强制转换")
        try:
            # 尝试转换为列表，如果可能的话
            metadata = list(metadata)
        except:
            logger.error("保存报告元数据: metadata无法转换为列表，将使用空列表")
            metadata = []
    
    try:
        # 确保目录存在
        dir_path = os.path.dirname(REPORT_METADATA_PATH)
        if not os.path.exists(dir_path):
            try:
                os.makedirs(dir_path, exist_ok=True)
                logger.info(f"创建报告元数据目录: {dir_path}")
            except Exception as mkdir_err:
                logger.error(f"创建报告元数据目录失败: {str(mkdir_err)}")
                return False
        
        # 先将数据写入临时文件，然后重命名，避免写入中断导致数据丢失
        temp_path = f"{REPORT_METADATA_PATH}.tmp"
        try:
            with open(temp_path, 'w', encoding='utf-8') as f:
                json.dump(metadata, f, ensure_ascii=False, indent=2)
                
            # 如果原文件存在，先创建备份
            if os.path.exists(REPORT_METADATA_PATH):
                backup_path = f"{REPORT_METADATA_PATH}.bak"
                try:
                    os.replace(REPORT_METADATA_PATH, backup_path)
                    logger.debug(f"创建元数据文件备份: {backup_path}")
                except Exception as backup_err:
                    logger.warning(f"创建元数据备份失败，继续保存: {str(backup_err)}")
            
            # 重命名临时文件为正式文件
            os.replace(temp_path, REPORT_METADATA_PATH)
            logger.info(f"成功保存报告元数据: {len(metadata)}条记录")
            return True
                
        except Exception as write_err:
            logger.error(f"写入报告元数据失败: {str(write_err)}")
            # 尝试删除临时文件
            if os.path.exists(temp_path):
                try:
                    os.remove(temp_path)
                except:
                    pass
            return False
    except Exception as e:
        logger.error(f"保存报告元数据失败: {str(e)}\n{traceback.format_exc()}")
        return False

# 添加报告记录到元数据
def add_report_record(report_info):
    if not report_info or not isinstance(report_info, dict):
        logger.error("添加报告记录失败: report_info参数无效")
        return False
        
    try:
        # 加载现有元数据
        metadata = load_report_metadata()
        
        # 确保metadata是列表
        if metadata is None:
            metadata = []
        elif not isinstance(metadata, list):
            logger.warning("加载的元数据不是列表格式，将重置为空列表")
            metadata = []
            
        # 添加新记录
        metadata.append(report_info)
        
        # 保存更新后的元数据
        try:
            result = save_report_metadata(metadata)
            logger.info(f"添加报告记录: {report_info.get('id')} - 结果: {'成功' if result else '失败'}")
            return result
        except Exception as save_err:
            logger.error(f"保存报告元数据失败: {str(save_err)}")
            return False
    except Exception as e:
        logger.error(f"添加报告记录失败: {str(e)}\n{traceback.format_exc()}")
        return False

# API路由 - 获取报告历史
@app.route('/api/reports/history', methods=['GET'])
def get_report_history():
    try:
        # 从文件加载报告元数据
        reports = load_report_metadata()
        
        # 确保reports是可迭代的，如果为None则返回空列表
        if reports is None:
            logger.warning("报告元数据加载返回None，将使用空列表")
            reports = []
            
        # 确保返回的是列表
        if not isinstance(reports, list):
            logger.warning(f"报告元数据不是列表格式: {type(reports)}，将转换为列表")
            # 如果是字典，且可能有items方法，尝试转换为列表
            if isinstance(reports, dict):
                try:
                    reports = list(reports.values())
                except:
                    reports = []
            else:
                try:
                    reports = list(reports)
                except:
                    reports = []
            
        # 按生成时间降序排序
        try:
            if reports and len(reports) > 0:
                reports.sort(key=lambda x: x.get('created_at', ''), reverse=True)
        except Exception as sort_err:
            logger.error(f"报告排序失败: {str(sort_err)}")
            # 排序失败不影响返回结果
        
        # 返回成功响应
        return jsonify(reports)
    except Exception as e:
        logger.error(f"获取报告历史失败: {str(e)}\n{traceback.format_exc()}")
        return handle_error(f"获取报告历史失败: {str(e)}")

# API路由 - 生成报告
@app.route('/api/reports/generate', methods=['POST'])
def generate_report():
    try:
        data = request.json
        if not data:
            return handle_error("请求数据为空", 400)
        
        # 提取参数
        report_type = data.get('type', 'daily')
        start_date = data.get('startDate')
        end_date = data.get('endDate')
        region = data.get('region', 'all')
        report_format = data.get('format', 'pdf').lower()  # 确保格式小写
        report_name = data.get('name', f"空气质量{get_report_type_name(report_type)}")
        
        # 记录请求信息
        logger.info(f"收到报告生成请求: 类型={report_type}, 区域={region}, 日期={start_date}至{end_date}, 格式={report_format}")
        
        # 基本参数检查
        if not start_date or not end_date:
            return handle_error("开始日期和结束日期不能为空", 400)
            
        # 生成报告ID - 确保为纯字母数字
        report_id = f"report_{datetime.now().strftime('%Y%m%d%H%M%S')}_{uuid.uuid4().hex[:8]}"
        
        # 确定输出目录
        if report_format == 'pdf':
            output_dir = REPORTS_PDF_DIR
            file_ext = '.pdf'
        elif report_format == 'excel':
            output_dir = REPORTS_EXCEL_DIR
            file_ext = '.xlsx'
        elif report_format == 'word':
            output_dir = REPORTS_WORD_DIR
            file_ext = '.docx'
        elif report_format == 'html':
            output_dir = REPORTS_HTML_DIR
            file_ext = '.html'
        else:
            output_dir = REPORTS_DIR
            file_ext = f".{report_format}"
            
        # 确保输出目录存在
        try:
            os.makedirs(output_dir, exist_ok=True)
            logger.info(f"确保报告输出目录存在: {output_dir}")
        except Exception as e:
            logger.error(f"创建输出目录失败: {str(e)}")
            return handle_error(f"创建输出目录失败: {str(e)}")
        
        # 构建输出文件路径
        output_filename = f"{report_id}{file_ext}"
        output_path = os.path.join(output_dir, output_filename)
        logger.info(f"报告将保存至: {output_path}")
        
        # 提取内容选项 - 根据前端传递的值映射到后端处理的字段
        content_options = {
            'overview': data.get('content', {}).get('overview', True), 
            'pollution': data.get('content', {}).get('pollution', True),
            'trend': data.get('content', {}).get('trend', True),
            'warning': data.get('content', {}).get('warning', True),
            'policy': data.get('content', {}).get('policy', False)
        }
        
        # 生成报告文件
        result = generate_report_content(
            report_type, start_date, end_date, 
            region, content_options, report_format, 
            output_path
        )
        
        if isinstance(result, dict) and 'error' in result:
            logger.error(f"报告生成失败: {result['error']}")
            return handle_error(result['error'])
            
        # 检查文件是否实际生成
        if not os.path.exists(output_path):
            error_msg = f"报告文件未能成功生成: {output_path}"
            logger.error(error_msg)
            return handle_error(error_msg)
            
        # 构建下载URL - 确保包含完整的文件扩展名
        download_url = f"/api/reports/download/{output_filename}"
        
        # 创建报告记录 - 添加报告名称和状态字段
        report_record = {
            'id': report_id,
            'name': report_name,
            'type': report_type,
            'format': report_format,
            'region': region,
            'start_date': start_date,
            'end_date': end_date,
            'created_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'file_path': output_path,
            'download_url': download_url,
            'status': 'completed',
            'file_name': output_filename
        }
        
        # 添加报告记录到元数据
        try:
            add_success = add_report_record(report_record)
            if not add_success:
                logger.warning(f"报告元数据添加失败，但报告文件已生成: {output_path}")
        except Exception as e:
            logger.error(f"添加报告记录到元数据时发生错误: {str(e)}")
            # 即使记录添加失败，也继续返回成功，因为文件已生成
        
        # 返回成功响应 - 包含完整的报告信息
        response = {
            'success': True,
            'id': report_id,
            'name': report_name,
            'type': report_type,
            'format': report_format,
            'download_url': download_url,
            'file_name': output_filename,
            'created_at': report_record['created_at'],
            'status': 'completed'
        }
        logger.info(f"报告生成成功: {report_id}, 文件路径: {output_path}")
        return jsonify(response)
            
    except Exception as e:
        logger.error(f"报告生成失败: {str(e)}\n{traceback.format_exc()}")
        return handle_error(f"报告生成失败: {str(e)}")

# API路由 - 下载报告
@app.route('/api/reports/download/<filename>', methods=['GET'])
def download_report(filename):
    try:
        if not filename:
            logger.error("下载请求缺少文件名")
            return handle_error("下载失败: 缺少报告ID", 400)
            
        # 记录下载请求
        logger.info(f"收到报告下载请求: {filename}")
        
        # 从元数据中查找报告
        reports = load_report_metadata()
        report = None
        
        # 提取report_id和扩展名
        report_id = filename.split('.')[0] if '.' in filename else filename
        ext = filename.split('.')[-1] if '.' in filename else ''
        
        # 先尝试通过文件名完全匹配
        if reports and isinstance(reports, list):
            for r in reports:
                if not isinstance(r, dict):
                    continue
                if r.get('file_name') == filename:
                    report = r
                    logger.debug(f"从元数据中找到报告文件: {filename}")
                    break
                if r.get('id') == report_id:
                    report = r
                    logger.debug(f"从元数据中通过ID找到报告: {report_id}")
                    break
        
        # 如果没找到，尝试在各个目录中查找文件
        if not report:
            # 根据扩展名确定可能的目录
            if ext == 'pdf':
                primary_dir = REPORTS_PDF_DIR
            elif ext in ['xlsx', 'xls']:
                primary_dir = REPORTS_EXCEL_DIR
            elif ext in ['docx', 'doc']:
                primary_dir = REPORTS_WORD_DIR
            elif ext == 'html':
                primary_dir = REPORTS_HTML_DIR
            else:
                primary_dir = REPORTS_DIR
                
            # 检查主要可能的位置
            primary_path = os.path.join(primary_dir, filename)
            if os.path.exists(primary_path) and os.path.isfile(primary_path):
                logger.info(f"找到报告文件: {primary_path}")
                # 创建临时报告记录
                report = {
                    'file_path': primary_path,
                    'file_name': filename
                }
            else:
                # 检查所有报告目录
                potential_paths = [
                    os.path.join(REPORTS_PDF_DIR, filename),
                    os.path.join(REPORTS_EXCEL_DIR, filename),
                    os.path.join(REPORTS_WORD_DIR, filename),
                    os.path.join(REPORTS_HTML_DIR, filename),
                    os.path.join(REPORTS_DIR, filename)
                ]
                
                for path in potential_paths:
                    if os.path.exists(path) and os.path.isfile(path):
                        logger.info(f"找到报告文件: {path}")
                        # 创建临时报告记录
                        report = {
                            'file_path': path,
                            'file_name': filename
                        }
                        break
        
        if not report:
            logger.error(f"报告文件不存在: {filename}")
            return handle_error("下载失败: 报告不存在或已被删除", 404)
        
        file_path = report.get('file_path', '')
        if not file_path:
            logger.error(f"报告记录缺少文件路径: {report}")
            return handle_error("报告记录缺少文件路径", 500)
        
        if not os.path.exists(file_path):
            logger.error(f"报告文件路径不存在: {file_path}")
            return handle_error(f"报告文件不存在: {file_path}", 404)
            
        if not os.path.isfile(file_path):
            logger.error(f"指定的路径不是文件: {file_path}")
            return handle_error(f"指定的路径不是文件: {file_path}", 400)
        
        # 获取文件MIME类型
        try:
            content_type, _ = mimetypes.guess_type(file_path)
            if not content_type:
                if file_path.endswith('.pdf'):
                    content_type = 'application/pdf'
                elif file_path.endswith('.xlsx'):
                    content_type = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
                elif file_path.endswith('.docx'):
                    content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                elif file_path.endswith('.html'):
                    content_type = 'text/html'
                else:
                    content_type = 'application/octet-stream'
        except Exception as mime_err:
            logger.warning(f"无法确定MIME类型: {str(mime_err)}")
            content_type = 'application/octet-stream'
        
        try:
            logger.info(f"发送报告文件: {file_path}, 类型: {content_type}")
            return send_file(
                file_path,
                mimetype=content_type,
                as_attachment=True,
                download_name=report.get('file_name', os.path.basename(file_path))
            )
        except Exception as send_err:
            logger.error(f"发送文件失败: {str(send_err)}")
            return handle_error(f"发送文件失败: {str(send_err)}", 500)
    except Exception as e:
        logger.error(f"下载报告失败: {str(e)}\n{traceback.format_exc()}")
        return handle_error(f"下载报告失败: {str(e)}")

# API路由 - 服务状态检查
@app.route('/api/health', methods=['GET'])
def health_check():
    """简单的健康检查API，返回服务状态"""
    return jsonify({
        'status': 'ok',
        'service': 'reports-api',
        'version': '1.0.0',
        'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    })

# API路由 - 服务状态
@app.route('/api/status', methods=['GET'])
def service_status():
    """详细的服务状态API，包括数据库连接状态、文件系统状态等"""
    try:
        # 检查数据库连接
        db_status = "ok"
        db_details = {}
        try:
            start_time = time.time()
            conn = mysql.connector.connect(**DB_CONFIG)
            # 执行简单的查询测试
            cursor = conn.cursor()
            cursor.execute("SELECT 1")
            cursor.fetchall()
            cursor.close()
            conn.close()
            db_details["response_time"] = f"{(time.time() - start_time):.3f}s"
        except Exception as e:
            db_status = "error"
            db_details["error"] = str(e)
        
        # 检查报告目录
        fs_status = {}
        for name, path in [
            ("reports_dir", REPORTS_DIR),
            ("pdf_dir", REPORTS_PDF_DIR),
            ("excel_dir", REPORTS_EXCEL_DIR),
            ("word_dir", REPORTS_WORD_DIR),
            ("html_dir", REPORTS_HTML_DIR),
            ("fonts_dir", FONTS_DIR)
        ]:
            if os.path.exists(path):
                fs_status[name] = "ok"
                try:
                    # 测试目录可写
                    test_file = os.path.join(path, ".test_write")
                    with open(test_file, 'w') as f:
                        f.write("test")
                    os.remove(test_file)
                except Exception as e:
                    fs_status[name] = f"not writable: {str(e)}"
            else:
                fs_status[name] = "missing"
        
        # 检查字体文件
        font_status = "ok" if os.path.exists(DEFAULT_FONT_PATH) else "missing"
        
        # 检查报告元数据
        metadata_status = "ok"
        try:
            reports = load_report_metadata()
            metadata_details = {"count": len(reports) if reports else 0}
        except Exception as e:
            metadata_status = "error"
            metadata_details = {"error": str(e)}
        
        # 构建完整状态报告
        status_report = {
            'status': 'ok',  # 默认值
            'service': 'reports-api',
            'components': {
                'database': {
                    'status': db_status,
                    'details': db_details
                },
                'filesystem': {
                    'status': 'ok' if all(s == "ok" for s in fs_status.values()) else 'warning',
                    'details': fs_status
                },
                'fonts': {
                    'status': font_status,
                    'path': DEFAULT_FONT_PATH
                },
                'metadata': {
                    'status': metadata_status,
                    'details': metadata_details
                }
            },
            'version': '1.0.0',
            'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        }
        
        # 检查是否有错误组件，更新整体状态
        has_error = False
        for component in status_report['components'].values():
            if component.get('status') == 'error':
                has_error = True
                break
        
        if has_error:
            status_report['status'] = 'error'
        
        return jsonify(status_report)
    except Exception as e:
        logger.error(f"获取服务状态失败: {str(e)}\n{traceback.format_exc()}")
        return handle_error(f"获取服务状态失败: {str(e)}")

# 根据AQI值获取空气质量等级
def get_air_quality_category(aqi):
    if aqi <= 50:
        return "优"
    elif aqi <= 100:
        return "良"
    elif aqi <= 150:
        return "轻度污染"
    elif aqi <= 200:
        return "中度污染"
    elif aqi <= 300:
        return "重度污染"
    else:
        return "严重污染"

# 生成报告内容
def generate_report_content(report_type, start_date, end_date, region, content_options, report_format, output_path):
    # 加载数据
    try:
        # 加载空气质量数据 - 确保禁用模拟数据
        data = load_air_quality_data(start_date, end_date, region)
        
        # 检查数据是否包含错误
        if 'error' in data:
            error_msg = data.get('error', '未知错误')
            logger.warning(f"数据加载过程中发生错误: {error_msg}")
            return {'error': f"数据加载失败: {error_msg}"}
        
        # 检查数据是否都为空 - 只有当AQI数据和污染物数据都为空时才返回错误
        aqi_empty = data.get('aqi_data') is None or data.get('aqi_data').empty
        poll_empty = data.get('pollutants') is None or data.get('pollutants').empty
        
        if aqi_empty and poll_empty:
            logger.warning(f"所选时间段({start_date}至{end_date})和地区({region})没有可用的空气质量数据")
            return {'error': f"所选时间段({start_date}至{end_date})和地区({get_region_name(region)})没有可用的空气质量数据，请选择其他时间范围或地区"}
        
        # 展示数据信息
        logger.info(f"生成报告 - AQI数据: {0 if aqi_empty else len(data['aqi_data'])}条, 污染物数据: {0 if poll_empty else len(data['pollutants'])}条")
        
        # 生成报告
        logger.info(f"生成{report_type}报告: {start_date}至{end_date}, 区域={region}, 格式={report_format}")
        logger.info(f"内容选项: {content_options}")
        
        # 将前端内容选项映射到后端内容选项
        mapped_options = {
            'air_quality': content_options.get('overview', True),
            'pollutants': content_options.get('pollution', True),
            'trends': content_options.get('trend', True),
            'alerts': content_options.get('warning', True),
            'recommendations': content_options.get('policy', False)
        }
        
        # 根据格式生成不同类型的报告
        if report_format == 'pdf':
            result = generate_pdf_report(data, mapped_options, output_path, report_type, region, start_date, end_date)
        elif report_format == 'excel':
            result = generate_excel_report(data, mapped_options, output_path, report_type, region, start_date, end_date)
        elif report_format == 'word':
            result = generate_word_report(data, mapped_options, output_path, report_type, region, start_date, end_date)
        elif report_format == 'html':
            result = generate_html_report(data, mapped_options, output_path, report_type, region, start_date, end_date)
        else:
            logger.error(f"不支持的报告格式: {report_format}")
            return {'error': f"不支持的报告格式: {report_format}"}
        
        # 检查报告文件是否成功生成
        if not os.path.exists(output_path):
            logger.error(f"报告文件未生成: {output_path}")
            return {'error': f"报告文件未生成: {output_path}"}
        
        # 检查文件大小是否合理（至少应该有1KB）
        file_size = os.path.getsize(output_path)
        if file_size < 1024:
            logger.warning(f"报告文件大小异常: {file_size} 字节")
        
        logger.info(f"报告生成成功: {output_path}, 大小: {file_size} 字节")
        return output_path
    except Exception as e:
        logger.error(f"报告生成失败: {str(e)}\n{traceback.format_exc()}")
        return {'error': f"报告生成失败: {str(e)}"}

# 生成PDF报告
def generate_pdf_report(data, content_options, output_path, report_type, region, start_date, end_date):
    """生成PDF格式的报告
    
    Args:
        data (dict): 数据字典，包含AQI数据和污染物数据
        content_options (dict): 内容选项，指定要包含的内容
        output_path (str): 输出文件路径
        report_type (str): 报告类型
        region (str): 地区
        start_date (str): 开始日期
        end_date (str): 结束日期
        
    Returns:
        str: 生成的报告文件路径，或者带有错误信息的字典
    """
    try:
        logger.info(f"开始生成PDF报告: {output_path}")
        logger.info(f"内容选项: {content_options}")
        
        # 检查数据有效性
        if data is None or not isinstance(data, dict):
            error_msg = "无效的数据格式"
            logger.error(error_msg)
            return {'error': error_msg}
            
        aqi_data = data.get('aqi_data')
        pollutants = data.get('pollutants')
        
        if aqi_data is None or pollutants is None:
            error_msg = "数据中缺少AQI数据或污染物数据"
            logger.error(error_msg)
            return {'error': error_msg}
            
        # 确保输出目录存在
        output_dir = os.path.dirname(output_path)
        os.makedirs(output_dir, exist_ok=True)
        
        # 注册中文字体
        font_path = DEFAULT_FONT_PATH
        if not os.path.exists(font_path):
            font_path = os.path.join(FONTS_DIR, 'simhei.ttf')
            if not os.path.exists(font_path):
                # 尝试使用系统字体
                system_fonts = ['C:/Windows/Fonts/simhei.ttf', '/usr/share/fonts/truetype/arphic/uming.ttc']
                for sf in system_fonts:
                    if os.path.exists(sf):
                        font_path = sf
                        break
        
        if os.path.exists(font_path):
            logger.info(f"使用字体: {font_path}")
            pdfmetrics.registerFont(TTFont('SimHei', font_path))
        else:
            logger.warning("找不到中文字体，将使用默认字体")
        
        # 创建PDF文档
        logger.info("创建PDF文档...")
        pdf_buffer = io.BytesIO()
        c = canvas.Canvas(pdf_buffer, pagesize=A4)
        width, height = A4
        
        # 页码计数器
        page_number = 1
        
        # 函数：添加页眉和页码
        def add_header_and_footer():
            # 页眉
            c.setFont('SimHei' if os.path.exists(font_path) else 'Helvetica', 10)
            c.drawString(50, height - 20, f"广东省空气质量监测系统 - {get_report_type_name(report_type)}报告")
            c.drawRightString(width - 50, height - 20, f"时间范围: {start_date} 至 {end_date}")
            
            # 页脚
            c.drawString(50, 20, f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            c.drawRightString(width - 50, 20, f"第 {page_number} 页")
        
        # 设置标题
        c.setFont('SimHei' if os.path.exists(font_path) else 'Helvetica', 18)
        title = f"广东省空气质量监测系统 - {get_report_type_name(report_type)}报告"
        c.drawCentredString(width/2, height-50, title)
        
        # 设置副标题
        c.setFont('SimHei' if os.path.exists(font_path) else 'Helvetica', 12)
        subtitle = f"时间范围: {start_date} 至 {end_date}   地区: {get_region_name(region)}"
        c.drawCentredString(width/2, height-80, subtitle)
        
        # 绘制分隔线
        c.line(50, height-90, width-50, height-90)
        
        # 添加页眉页脚
        add_header_and_footer()
        
        # 当前位置（Y坐标）
        y_position = height - 120
        
        # 添加摘要
        if not aqi_data.empty:
            c.setFont('SimHei' if os.path.exists(font_path) else 'Helvetica', 12)
            c.drawString(50, y_position, "空气质量摘要")
            y_position -= 20
            
            # 计算平均AQI
            avg_aqi = round(aqi_data['aqi'].mean(), 1)
            # 获取最高AQI
            max_aqi = aqi_data['aqi'].max()
            # 获取最高AQI的城市
            max_aqi_city = aqi_data.loc[aqi_data['aqi'].idxmax()]['city'] if len(aqi_data) > 0 else "无数据"
            
            # 绘制摘要信息
            c.setFont('SimHei' if os.path.exists(font_path) else 'Helvetica', 10)
            c.drawString(70, y_position, f"平均空气质量指数(AQI): {avg_aqi}")
            y_position -= 15
            c.drawString(70, y_position, f"最高空气质量指数: {max_aqi} ({max_aqi_city})")
            y_position -= 15
            c.drawString(70, y_position, f"空气质量等级: {get_air_quality_category(avg_aqi)}")
            y_position -= 30
        
        # 添加空气质量概览部分
        if content_options.get('air_quality', True) and not aqi_data.empty:
            logger.info("添加空气质量概览...")
            c.setFont('SimHei' if os.path.exists(font_path) else 'Helvetica', 12)
            c.drawString(50, y_position, "空气质量概览")
            y_position -= 20
            
            # 创建空气质量数据表格
            cities = aqi_data['city'].unique()
            rows = [["城市", "AQI", "质量等级"]]
            
            # 统计城市数量，以便规划表格布局
            total_cities = len(cities)
            logger.info(f"报告将包含 {total_cities} 个城市的数据")
            
            # 为所有城市添加数据
            for city in cities:
                city_data = aqi_data[aqi_data['city'] == city]
                avg_city_aqi = round(city_data['aqi'].mean(), 1)
                quality_level = get_air_quality_category(avg_city_aqi)
                rows.append([city, str(avg_city_aqi), quality_level])
            
            # 绘制表格
            table_width = width - 100
            col_widths = [table_width * 0.4, table_width * 0.3, table_width * 0.3]
            row_height = 20
            rows_per_page = 25  # 每页最多显示的行数
            
            # 计算表格需要多少页
            total_rows = len(rows)
            
            # 如果数据太多，需要分页
            if y_position - (total_rows * row_height) < 50:
                # 第一页能容纳的行数（考虑已用空间）
                first_page_rows = max(1, min(rows_per_page, int((y_position - 50) / row_height)))
                
                # 绘制表头和第一页数据
                for i in range(min(first_page_rows, total_rows)):
                    row = rows[i]
                    for j, cell in enumerate(row):
                        x = 50 + sum(col_widths[:j])
                        y = y_position - i * row_height
                        
                        # 绘制单元格
                        if i == 0:  # 表头
                            c.setFillColorRGB(0.9, 0.9, 0.9)
                            c.rect(x, y - 15, col_widths[j], row_height, stroke=1, fill=1)
                            c.setFillColorRGB(0, 0, 0)  # 重置填充颜色
                            c.setFont('SimHei' if os.path.exists(font_path) else 'Helvetica-Bold', 9)
                            # 居中显示表头文字
                            text_width = c.stringWidth(cell, 'SimHei' if os.path.exists(font_path) else 'Helvetica-Bold', 9)
                            text_x = x + (col_widths[j] - text_width) / 2
                            c.drawString(text_x, y - 5, cell)
                        else:
                            c.setFont('SimHei' if os.path.exists(font_path) else 'Helvetica', 9)
                            # 居中显示数据
                            text_width = c.stringWidth(cell, 'SimHei' if os.path.exists(font_path) else 'Helvetica', 9)
                            text_x = x + (col_widths[j] - text_width) / 2
                            c.drawString(text_x, y - 5, cell)
                            c.rect(x, y - 15, col_widths[j], row_height, stroke=1, fill=0)
                
                # 处理剩余的行
                if total_rows > first_page_rows:
                    # 开始新页面
                    c.showPage()
                    page_number += 1
                    add_header_and_footer()
                    
                    # 重新绘制表头
                    y_position = height - 120
                    c.setFont('SimHei' if os.path.exists(font_path) else 'Helvetica', 12)
                    c.drawString(50, y_position, "空气质量概览 (续)")
                    y_position -= 20
                    
                    # 在新页面重新绘制表头
                    header_row = rows[0]
                    for j, cell in enumerate(header_row):
                        x = 50 + sum(col_widths[:j])
                        y = y_position
                        
                        c.setFillColorRGB(0.9, 0.9, 0.9)
                        c.rect(x, y - 15, col_widths[j], row_height, stroke=1, fill=1)
                        c.setFillColorRGB(0, 0, 0)
                        c.setFont('SimHei' if os.path.exists(font_path) else 'Helvetica-Bold', 9)
                        c.drawString(x + 2, y, cell)
                        c.rect(x, y - 15, col_widths[j], row_height, stroke=1, fill=0)
                    
                    # 计算剩余的行
                    remaining_rows = total_rows - first_page_rows
                    rows_per_new_page = min(rows_per_page - 1, int((height - 150) / row_height))  # 减去表头和页眉页脚
                    pages_needed = (remaining_rows + rows_per_new_page - 1) // rows_per_new_page
                    
                    row_index = first_page_rows
                    for page in range(pages_needed):
                        start_row = row_index
                        end_row = min(row_index + rows_per_new_page, total_rows)
                        
                        for i in range(start_row, end_row):
                            row = rows[i]
                            row_pos = i - start_row + 1  # +1是因为表头占一行
                            
                            for j, cell in enumerate(row):
                                x = 50 + sum(col_widths[:j])
                                y = y_position - row_pos * row_height
                                
                                c.setFont('SimHei' if os.path.exists(font_path) else 'Helvetica', 9)
                                c.drawString(x + 2, y - 5, cell)
                                c.rect(x, y - 15, col_widths[j], row_height, stroke=1, fill=0)
                        
                        row_index = end_row
                        
                        # 如果还有更多页，创建新页面
                        if page < pages_needed - 1:
                            c.showPage()
                            page_number += 1
                            add_header_and_footer()
                            
                            # 重设Y位置
                            y_position = height - 120
                            c.setFont('SimHei' if os.path.exists(font_path) else 'Helvetica', 12)
                            c.drawString(50, y_position, "空气质量概览 (续)")
                            y_position -= 20
                            
                            # 重新绘制表头
                            header_row = rows[0]
                            for j, cell in enumerate(header_row):
                                x = 50 + sum(col_widths[:j])
                                y = y_position
                                
                                c.setFillColorRGB(0.9, 0.9, 0.9)
                                c.rect(x, y - 15, col_widths[j], row_height, stroke=1, fill=1)
                                c.setFillColorRGB(0, 0, 0)
                                c.setFont('SimHei' if os.path.exists(font_path) else 'Helvetica-Bold', 9)
                                c.drawString(x + 2, y, cell)
                                c.rect(x, y - 15, col_widths[j], row_height, stroke=1, fill=0)
            else:
                # 如果一页足够，直接绘制所有行
                for i, row in enumerate(rows):
                    for j, cell in enumerate(row):
                        x = 50 + sum(col_widths[:j])
                        y = y_position - i * row_height
                        
                        if i == 0:  # 表头
                            c.setFillColorRGB(0.9, 0.9, 0.9)
                            c.rect(x, y - 15, col_widths[j], row_height, stroke=1, fill=1)
                            c.setFillColorRGB(0, 0, 0)
                            c.setFont('SimHei' if os.path.exists(font_path) else 'Helvetica-Bold', 9)
                        else:
                            c.setFont('SimHei' if os.path.exists(font_path) else 'Helvetica', 9)
                            
                        c.drawString(x + 2, y - 5, cell)
                        c.rect(x, y - 15, col_widths[j], row_height, stroke=1, fill=0)
            
            # 更新Y位置（前进至表格末尾）
            if total_rows <= rows_per_page:
                y_position -= (total_rows + 1) * row_height
            else:
                # 如果已分页，则Y位置是当前页最后绘制的位置
                y_position -= ((row_index - start_row) + 2) * row_height
            
            y_position -= 20  # 与下一部分保持间距
        
        # 添加污染物分析部分
        if content_options.get('pollutants', True) and not pollutants.empty:
            # 如果空间不足，换页
            if y_position < 200:
                c.showPage()
                page_number += 1
                add_header_and_footer()
                y_position = height - 120
            
            logger.info("添加污染物分析...")
            c.setFont('SimHei' if os.path.exists(font_path) else 'Helvetica', 12)
            c.drawString(50, y_position, "污染物分析")
            y_position -= 25  # 适当间距
            
            # 计算各污染物平均值
            pollutants_avg = {
                'PM2.5': round(pollutants['pm25'].mean(), 1),
                'PM10': round(pollutants['pm10'].mean(), 1),
                'SO2': round(pollutants['so2'].mean(), 1),
                'NO2': round(pollutants['no2'].mean(), 1),
                'O3': round(pollutants['o3'].mean(), 1),
                'CO': round(pollutants['co'].mean(), 2)
            }
            
            # 绘制污染物均值表格
            c.setFont('SimHei' if os.path.exists(font_path) else 'Helvetica', 10)
            c.drawString(70, y_position, "主要污染物平均浓度:")
            y_position -= 25  # 适当间距
            
            # 绘制表头
            pollutants_table_width = width - 140
            col_width = pollutants_table_width / 6  # 6个污染物
            row_height = 30  # 大幅增加行高
            
            # 表头
            headers = ['PM2.5', 'PM10', 'SO2', 'NO2', 'O3', 'CO']
            for i, header in enumerate(headers):
                x = 70 + i * col_width
                c.setFillColorRGB(0.9, 0.9, 0.9)
                c.rect(x, y_position - 20, col_width, row_height, stroke=1, fill=1)
                c.setFillColorRGB(0, 0, 0)
                c.setFont('SimHei' if os.path.exists(font_path) else 'Helvetica-Bold', 9)
                # 居中显示表头文字
                text_width = c.stringWidth(header, 'SimHei' if os.path.exists(font_path) else 'Helvetica-Bold', 9)
                text_x = x + (col_width - text_width) / 2
                c.drawString(text_x, y_position - 5, header)
            
            y_position -= row_height + 5  # 增加行间距
            
            # 数据行
            values = list(pollutants_avg.values())
            for i, value in enumerate(values):
                x = 70 + i * col_width
                c.setFont('SimHei' if os.path.exists(font_path) else 'Helvetica', 9)
                c.rect(x, y_position - 20, col_width, row_height, stroke=1, fill=0)
                # 居中显示数据
                value_str = str(value)
                text_width = c.stringWidth(value_str, 'SimHei' if os.path.exists(font_path) else 'Helvetica', 9)
                text_x = x + (col_width - text_width) / 2
                c.drawString(text_x, y_position - 5, value_str)
            
            y_position -= (row_height + 40)  # 大幅增加间距
            
            # 添加各城市污染物表格
            c.setFont('SimHei' if os.path.exists(font_path) else 'Helvetica', 10)
            c.drawString(70, y_position, "各城市主要污染物指标:")
            y_position -= 25  # 适当间距
            
            # 创建城市污染物表格
            cities = pollutants['city'].unique()
            rows = [["城市", "PM2.5", "PM10", "SO2", "NO2", "O3", "CO"]]
            
            for city in cities:
                city_data = pollutants[pollutants['city'] == city]
                pm25 = round(city_data['pm25'].mean(), 1)
                pm10 = round(city_data['pm10'].mean(), 1)
                so2 = round(city_data['so2'].mean(), 1)
                no2 = round(city_data['no2'].mean(), 1)
                o3 = round(city_data['o3'].mean(), 1)
                co = round(city_data['co'].mean(), 2)
                
                rows.append([city, str(pm25), str(pm10), str(so2), str(no2), str(o3), str(co)])
            
            # 设置表格参数
            table_width = width - 140  # 表格总宽度
            col_widths = [
                table_width * 0.25,  # 城市名称列宽度
                table_width * 0.125,  # PM2.5列宽度
                table_width * 0.125,  # PM10列宽度
                table_width * 0.125,  # SO2列宽度
                table_width * 0.125,  # NO2列宽度
                table_width * 0.125,  # O3列宽度
                table_width * 0.125   # CO列宽度
            ]
            
            # 固定参数
            row_height = 25  # 行高
            x_start = 70     # 表格起始x坐标
            max_rows_per_page = 15  # 每页最多行数
            
            # 绘制表格
            row_index = 0
            while row_index < len(rows):
                # 检查是否需要新页面
                if row_index > 0 and (row_index % max_rows_per_page == 0 or y_position < 150):
                    c.showPage()
                    page_number += 1
                    add_header_and_footer()
                    y_position = height - 120
                    
                    c.setFont('SimHei' if os.path.exists(font_path) else 'Helvetica', 12)
                    c.drawString(50, y_position, "污染物分析 (续)")
                    y_position -= 25
                    
                    # 重新绘制表头
                    draw_table_header = True
                else:
                    draw_table_header = (row_index == 0)
                
                # 计算本页要绘制的行数
                remaining_rows = len(rows) - row_index
                rows_this_page = min(remaining_rows, max_rows_per_page if row_index > 0 else max_rows_per_page - 1)
                
                # 绘制表头
                if draw_table_header:
                    header_row = rows[0]
                    for i, header in enumerate(header_row):
                        x = x_start + sum(col_widths[:i])
                        
                        # 绘制表头背景和边框
                        c.setFillColorRGB(0.9, 0.9, 0.9)  # 浅灰色背景
                        c.rect(x, y_position - row_height, col_widths[i], row_height, stroke=1, fill=1)
                        c.setFillColorRGB(0, 0, 0)  # 重置为黑色
                        
                        # 绘制表头文字
                        c.setFont('SimHei' if os.path.exists(font_path) else 'Helvetica-Bold', 9)
                        text_width = c.stringWidth(header, 'SimHei' if os.path.exists(font_path) else 'Helvetica-Bold', 9)
                        text_x = x + (col_widths[i] - text_width) / 2  # 水平居中
                        c.drawString(text_x, y_position - row_height + 8, header)  # 垂直居中
                    
                    y_position -= row_height
                    
                    # 如果是第一页，跳过表头行，因为已经绘制了
                    if row_index == 0:
                        row_index += 1
                
                # 绘制数据行
                for i in range(rows_this_page):
                    current_row = row_index + i
                    if current_row >= len(rows):
                        break
                        
                    row = rows[current_row]
                    
                    for j, cell in enumerate(row):
                        x = x_start + sum(col_widths[:j])
                        
                        # 绘制单元格边框
                        c.rect(x, y_position - row_height, col_widths[j], row_height, stroke=1, fill=0)
                        
                        # 绘制单元格内容
                        c.setFont('SimHei' if os.path.exists(font_path) else 'Helvetica', 9)
                        
                        if j == 0:  # 城市名称列左对齐
                            c.drawString(x + 5, y_position - row_height + 8, cell)
                        else:  # 数值列居中对齐
                            text_width = c.stringWidth(cell, 'SimHei' if os.path.exists(font_path) else 'Helvetica', 9)
                            text_x = x + (col_widths[j] - text_width) / 2
                            c.drawString(text_x, y_position - row_height + 8, cell)
                    
                    y_position -= row_height
                
                # 更新行索引
                row_index += rows_this_page
                
                # 如果还有更多行要绘制，则添加新页面
                if row_index < len(rows):
                    continue
            
            # 表格结束后添加适当间距
            y_position -= 20
        
        # 添加趋势变化部分
        if content_options.get('trends', True) and not aqi_data.empty:
            # 如果空间不足，换页
            if y_position < 200:
                c.showPage()
                page_number += 1
                add_header_and_footer()
                y_position = height - 120
            
            logger.info("添加趋势变化...")
            c.setFont('SimHei' if os.path.exists(font_path) else 'Helvetica', 12)
            c.drawString(50, y_position, "趋势变化")
            y_position -= 20
            
            # 对于PDF，添加趋势说明文字
            c.setFont('SimHei' if os.path.exists(font_path) else 'Helvetica', 10)
            c.drawString(70, y_position, "空气质量指数变化趋势(PDF格式限制，无法显示图表):")
            y_position -= 15
            
            # 计算并描述变化趋势
            if len(aqi_data['date'].unique()) > 1:
                # 有多天数据，可计算趋势
                earliest_date = aqi_data['date'].min()
                latest_date = aqi_data['date'].max()
                earliest_avg = aqi_data[aqi_data['date'] == earliest_date]['aqi'].mean()
                latest_avg = aqi_data[aqi_data['date'] == latest_date]['aqi'].mean()
                
                change_pct = ((latest_avg - earliest_avg) / earliest_avg) * 100 if earliest_avg > 0 else 0
                
                trend_text = f"从 {earliest_date} 到 {latest_date}，平均AQI从 {earliest_avg:.1f} 变化到 {latest_avg:.1f}，"
                
                if change_pct > 5:
                    trend_text += f"上升了 {abs(change_pct):.1f}%，空气质量有所下降。"
                elif change_pct < -5:
                    trend_text += f"下降了 {abs(change_pct):.1f}%，空气质量有所改善。"
                else:
                    trend_text += "变化不大，空气质量基本稳定。"
                
                c.drawString(70, y_position, trend_text)
            else:
                c.drawString(70, y_position, "监测时间段过短，无法显示变化趋势。")
            
            y_position -= 30
        
        # 添加预警信息部分
        if content_options.get('alerts', True) and not aqi_data.empty:
            # 如果空间不足，换页
            if y_position < 200:
                c.showPage()
                page_number += 1
                add_header_and_footer()
                y_position = height - 120
            
            logger.info("添加预警信息...")
            c.setFont('SimHei' if os.path.exists(font_path) else 'Helvetica', 12)
            c.drawString(50, y_position, "预警信息")
            y_position -= 20
            
            # 查找AQI超标的城市
            high_aqi_data = aqi_data[aqi_data['aqi'] > 100]
            
            if not high_aqi_data.empty:
                c.setFont('SimHei' if os.path.exists(font_path) else 'Helvetica', 10)
                c.drawString(70, y_position, f"发现 {len(high_aqi_data)} 个城市AQI超过100，可能对敏感人群健康造成影响:")
                y_position -= 15
                
                # 列出超标城市
                cities_warning = high_aqi_data.groupby('city')['aqi'].mean().reset_index()
                cities_warning = cities_warning.sort_values(by='aqi', ascending=False)
                
                warning_text = ", ".join([f"{row['city']}({row['aqi']:.1f})" for _, row in cities_warning.iterrows()])
                
                # 处理长文本换行
                max_width = width - 140
                words = warning_text.split(", ")
                lines = []
                current_line = ""
                
                for word in words:
                    test_line = current_line + word + ", " if current_line else word + ", "
                    if c.stringWidth(test_line, 'SimHei' if os.path.exists(font_path) else 'Helvetica', 9) < max_width:
                        current_line = test_line
                    else:
                        lines.append(current_line)
                        current_line = word + ", "
                
                if current_line:
                    lines.append(current_line)
                
                # 绘制文本
                for line in lines:
                    c.drawString(90, y_position, line)
                    y_position -= 15
            else:
                c.setFont('SimHei' if os.path.exists(font_path) else 'Helvetica', 10)
                c.drawString(70, y_position, "监测期间内未发现明显空气质量超标情况。")
                y_position -= 15
            
            y_position -= 15
        
        # 添加政策建议部分
        if content_options.get('recommendations', False):
            # 如果空间不足，换页
            if y_position < 200:
                c.showPage()
                page_number += 1
                add_header_and_footer()
                y_position = height - 120
            
            logger.info("添加政策建议...")
            c.setFont('SimHei' if os.path.exists(font_path) else 'Helvetica', 12)
            c.drawString(50, y_position, "政策建议")
            y_position -= 20
            
            # 根据空气质量给出建议
            recommendations = [
                "实施工业企业限产停产措施，减少污染物排放",
                "加强道路扬尘治理，增加洒水降尘频次",
                "建议公众减少户外活动，外出佩戴口罩",
                "加大环保执法力度，严厉打击违法排污行为"
            ]
            
            # 绘制建议列表
            c.setFont('SimHei' if os.path.exists(font_path) else 'Helvetica', 10)
            for i, rec in enumerate(recommendations):
                c.drawString(70, y_position, "• " + rec)
                y_position -= 15
        
        # 完成文档
        c.save()
        
        # 获取PDF内容并写入文件
        pdf_data = pdf_buffer.getvalue()
        with open(output_path, 'wb') as f:
            f.write(pdf_data)
        
        # 验证文件是否已创建
        if os.path.exists(output_path):
            logger.info(f"PDF报告生成成功: {output_path}")
            return output_path
        else:
            error_msg = f"PDF文件写入失败: {output_path}"
            logger.error(error_msg)
            return {'error': error_msg}
    except Exception as e:
        error_msg = f"生成PDF报告失败: {str(e)}"
        logger.error(f"{error_msg}\n{traceback.format_exc()}")
        return {'error': error_msg}

# 生成Excel报告
def generate_excel_report(data, content_options, output_path, report_type, region, start_date, end_date):
    try:
        # 创建Excel工作簿
        wb = openpyxl.Workbook()
        
        # 创建报告信息工作表
        info_sheet = wb.active
        info_sheet.title = "报告信息"
        
        # 添加标题和基本信息
        info_sheet['A1'] = "广东省空气质量监测报告"
        info_sheet['A1'].font = Font(size=16, bold=True)
        info_sheet.merge_cells('A1:D1')
        info_sheet['A1'].alignment = Alignment(horizontal='center')
        
        info_sheet['A3'] = "区域:"
        info_sheet['B3'] = region if region != 'all' else '全省'
        info_sheet['A4'] = "时间范围:"
        info_sheet['B4'] = f"{start_date} 至 {end_date}"
        info_sheet['A5'] = "生成时间:"
        info_sheet['B5'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        
        # 1. 空气质量数据工作表
        if content_options.get('air_quality', True) and not data['aqi_data'].empty:
            aqi_sheet = wb.create_sheet("空气质量数据")
            
            # 添加表头
            headers = ['日期', 'AQI', '质量等级', '城市']
            for col, header in enumerate(headers, 1):
                aqi_sheet.cell(row=1, column=col, value=header).font = Font(bold=True)
            
            # 添加数据
            aqi_data = data['aqi_data']
            for row_idx, (_, row) in enumerate(aqi_data.iterrows(), 2):
                aqi_sheet.cell(row=row_idx, column=1, value=str(row.get('date', '')))
                aqi_sheet.cell(row=row_idx, column=2, value=row.get('aqi', ''))
                aqi_sheet.cell(row=row_idx, column=3, value=row.get('quality_level', ''))
                aqi_sheet.cell(row=row_idx, column=4, value=row.get('city', ''))
        
        # 2. 污染物数据工作表
        if content_options.get('pollutants', True) and not data['pollutants'].empty:
            poll_sheet = wb.create_sheet("污染物数据")
            
            # 添加表头
            headers = ['日期', 'PM2.5', 'PM10', 'SO2', 'NO2', 'O3', 'CO', '城市']
            for col, header in enumerate(headers, 1):
                poll_sheet.cell(row=1, column=col, value=header).font = Font(bold=True)
            
            # 添加数据
            poll_data = data['pollutants']
            for row_idx, (_, row) in enumerate(poll_data.iterrows(), 2):
                poll_sheet.cell(row=row_idx, column=1, value=str(row.get('date', '')))
                poll_sheet.cell(row=row_idx, column=2, value=row.get('pm25', ''))
                poll_sheet.cell(row=row_idx, column=3, value=row.get('pm10', ''))
                poll_sheet.cell(row=row_idx, column=4, value=row.get('so2', ''))
                poll_sheet.cell(row=row_idx, column=5, value=row.get('no2', ''))
                poll_sheet.cell(row=row_idx, column=6, value=row.get('o3', ''))
                poll_sheet.cell(row=row_idx, column=7, value=row.get('co', ''))
                poll_sheet.cell(row=row_idx, column=8, value=row.get('city', ''))
        
        # 3. 统计分析工作表
        stats_sheet = wb.create_sheet("统计分析")
        
        row_idx = 1
        stats_sheet.cell(row=row_idx, column=1, value="统计指标").font = Font(bold=True)
        stats_sheet.cell(row=row_idx, column=2, value="数值").font = Font(bold=True)
        row_idx += 1
        
        if not data['aqi_data'].empty:
            aqi_data = data['aqi_data']
            
            stats_sheet.cell(row=row_idx, column=1, value="AQI数据记录数")
            stats_sheet.cell(row=row_idx, column=2, value=len(aqi_data))
            row_idx += 1
            
            stats_sheet.cell(row=row_idx, column=1, value="平均AQI")
            stats_sheet.cell(row=row_idx, column=2, value=round(aqi_data['aqi'].mean(), 2))
            row_idx += 1
            
            stats_sheet.cell(row=row_idx, column=1, value="最高AQI")
            stats_sheet.cell(row=row_idx, column=2, value=round(aqi_data['aqi'].max(), 2))
            row_idx += 1
            
            stats_sheet.cell(row=row_idx, column=1, value="最低AQI")
            stats_sheet.cell(row=row_idx, column=2, value=round(aqi_data['aqi'].min(), 2))
            row_idx += 1
        
        if not data['pollutants'].empty:
            pollutants_df = data['pollutants']
            row_idx += 1
            stats_sheet.cell(row=row_idx, column=1, value="污染物平均值").font = Font(bold=True)
            row_idx += 1
            
            # 计算主要污染物平均值
            poll_list = [('PM2.5', 'pm25'), ('PM10', 'pm10'), ('SO2', 'so2'), 
                        ('NO2', 'no2'), ('O3', 'o3'), ('CO', 'co')]
            
            for display_name, col_name in poll_list:
                if col_name in pollutants_df.columns:
                    stats_sheet.cell(row=row_idx, column=1, value=display_name)
                    stats_sheet.cell(row=row_idx, column=2, value=round(pollutants_df[col_name].mean(), 2))
                    row_idx += 1
        
        # 自动调整列宽
        for sheet in wb.sheetnames:
            ws = wb[sheet]
            for column in ws.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = (max_length + 2)
                ws.column_dimensions[column_letter].width = adjusted_width
        
        # 保存Excel文件
        wb.save(output_path)
        return output_path
    except Exception as e:
        logger.error(f"Excel生成失败: {str(e)}\n{traceback.format_exc()}")
        return {'error': str(e)}

# 生成Word报告
def generate_word_report(data, content_options, output_path, report_type, region, start_date, end_date):
    try:
        # 创建Word文档
        doc = Document()
        
        # 设置中文字体（如果需要）
        try:
            # 设置文档默认字体
            doc.styles['Normal'].font.name = '宋体'
            doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        except:
            logger.warning("无法设置中文字体，将使用默认字体")
        
        # 添加标题
        doc.add_heading('广东省空气质量监测报告', level=0)
        
        # 添加基本信息
        doc.add_paragraph(f"区域: {region if region != 'all' else '全省'}")
        doc.add_paragraph(f"时间范围: {start_date} 至 {end_date}")
        doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        # 添加分隔线
        doc.add_paragraph('_' * 50)
        
        # 1. 空气质量概览
        if content_options.get('air_quality', True) and not data['aqi_data'].empty:
            doc.add_heading('1. 空气质量概览', level=1)
            aqi_data = data['aqi_data']
            
            # 计算统计信息
            avg_aqi = aqi_data['aqi'].mean()
            max_aqi = aqi_data['aqi'].max()
            min_aqi = aqi_data['aqi'].min()
            
            p = doc.add_paragraph()
            p.add_run(f"数据记录数: {len(aqi_data)}\n")
            p.add_run(f"平均AQI: {avg_aqi:.1f}\n")
            p.add_run(f"最高AQI: {max_aqi:.1f}\n")
            p.add_run(f"最低AQI: {min_aqi:.1f}\n")
            
            # 添加空气质量数据表格
            doc.add_heading('空气质量数据表', level=2)
            
            # 创建表格
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            
            # 添加表头
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '日期'
            hdr_cells[1].text = 'AQI'
            hdr_cells[2].text = '质量等级'
            hdr_cells[3].text = '城市'
            
            # 添加数据行 (限制为前10行以避免文档过长)
            for _, row in aqi_data.head(10).iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = str(row.get('date', ''))
                row_cells[1].text = str(row.get('aqi', ''))
                row_cells[2].text = str(row.get('quality_level', ''))
                row_cells[3].text = str(row.get('city', ''))
        
        # 2. 污染物分析
        if content_options.get('pollutants', True) and not data['pollutants'].empty:
            doc.add_heading('2. 主要污染物分析', level=1)
            
            pollutants_df = data['pollutants']
            
            # 计算主要污染物平均值
            pollutants_means = {
                'PM2.5': pollutants_df['pm25'].mean(),
                'PM10': pollutants_df['pm10'].mean(),
                'SO2': pollutants_df['so2'].mean(),
                'NO2': pollutants_df['no2'].mean(),
                'O3': pollutants_df['o3'].mean(),
                'CO': pollutants_df['co'].mean()
            }
            
            p = doc.add_paragraph()
            for poll, value in pollutants_means.items():
                p.add_run(f"{poll}: {value:.1f}\n")
        
        # 3. 趋势分析
        if content_options.get('trends', True):
            doc.add_heading('3. 空气质量趋势', level=1)
            doc.add_paragraph("本报告Word版不包含图表，请选择HTML或Excel格式查看完整趋势图。")
        
        # 4. 预警信息
        if content_options.get('alerts', True):
            doc.add_heading('4. 空气质量预警信息', level=1)
            
            # 检查是否有超标情况
            if not data['aqi_data'].empty:
                high_aqi = data['aqi_data'][data['aqi_data']['aqi'] > 100]
                if not high_aqi.empty:
                    doc.add_paragraph(f"发现 {len(high_aqi)} 天空气质量超过100，可能对敏感人群健康造成影响。")
                else:
                    doc.add_paragraph("监测期间内未发现明显空气质量超标情况。")
        
        # 5. 政策建议
        if content_options.get('recommendations', False):
            doc.add_heading('5. 改善空气质量的建议', level=1)
            
            recommendations = [
                "加强工业排放监管，严格执行排放标准",
                "推广清洁能源使用，减少化石燃料消耗",
                "加强城市绿化建设，增加空气净化能力",
                "提高公众环保意识，鼓励绿色出行方式"
            ]
            
            for rec in recommendations:
                doc.add_paragraph(rec, style='List Bullet')
        
        # 保存Word文档
        doc.save(output_path)
        return output_path
    except Exception as e:
        logger.error(f"Word生成失败: {str(e)}\n{traceback.format_exc()}")
        return {'error': str(e)}

# 生成HTML报告
def generate_html_report(data, content_options, output_path, report_type, region, start_date, end_date):
    try:
        # 计算统计信息
        aqi_stats = {}
        pollutants_stats = {}
        
        if not data['aqi_data'].empty:
            aqi_data = data['aqi_data']
            aqi_stats = {
                'count': len(aqi_data),
                'avg': round(aqi_data['aqi'].mean(), 1),
                'max': round(aqi_data['aqi'].max(), 1),
                'min': round(aqi_data['aqi'].min(), 1),
            }
        
        if not data['pollutants'].empty:
            pollutants_df = data['pollutants']
            pollutants_stats = {
                'PM2.5': round(pollutants_df['pm25'].mean(), 1),
                'PM10': round(pollutants_df['pm10'].mean(), 1),
                'SO2': round(pollutants_df['so2'].mean(), 1),
                'NO2': round(pollutants_df['no2'].mean(), 1),
                'O3': round(pollutants_df['o3'].mean(), 1),
                'CO': round(pollutants_df['co'].mean(), 1)
            }
        
        # 创建HTML内容 - 头部
        html = f"""<!DOCTYPE html>
<html lang="zh-CN">
        <head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>广东省空气质量监测报告</title>
            <style>
        body {{ 
            font-family: "Microsoft YaHei", Arial, sans-serif; 
            margin: 0;
            padding: 20px;
            color: #333;
        }}
        .container {{
            max-width: 1000px;
            margin: 0 auto;
            background-color: #fff;
            padding: 20px;
            box-shadow: 0 0 10px rgba(0,0,0,0.1);
        }}
        h1, h2, h3 {{ color: #0066cc; }}
        h1 {{ text-align: center; margin-bottom: 30px; }}
        .header-info {{ 
            text-align: center; 
            margin-bottom: 30px;
            color: #666;
        }}
        .timestamp {{
            text-align: right;
            color: #888;
            font-size: 0.9em;
            margin-bottom: 20px;
        }}
        table {{ 
            border-collapse: collapse; 
            width: 100%; 
            margin-bottom: 20px;
        }}
        th, td {{ 
            border: 1px solid #ddd; 
            padding: 10px; 
            text-align: center;
        }}
        th {{ 
            background-color: #f2f2f2; 
            font-weight: bold;
        }}
        tr:nth-child(even) {{ background-color: #f9f9f9; }}
        .stats-card {{
            display: inline-block;
            background-color: #f8f9fa;
            border-radius: 5px;
            padding: 15px;
            margin: 10px;
            min-width: 120px;
            text-align: center;
            box-shadow: 0 2px 5px rgba(0,0,0,0.1);
        }}
        .stats-card .value {{
            font-size: 24px;
            font-weight: bold;
            color: #0066cc;
            margin: 10px 0;
        }}
        .stats-card .label {{
            font-size: 14px;
            color: #666;
        }}
        .stats-container {{
            display: flex;
            flex-wrap: wrap;
            justify-content: center;
            margin-bottom: 30px;
        }}
        .section {{
            margin-bottom: 40px;
            border-bottom: 1px solid #eee;
            padding-bottom: 20px;
        }}
        .pollutant-grid {{
            display: grid;
            grid-template-columns: repeat(3, 1fr);
            gap: 15px;
            margin-top: 20px;
        }}
        .pollutant-card {{
            background-color: #f8f9fa;
            border-radius: 5px;
            padding: 15px;
            text-align: center;
            box-shadow: 0 2px 5px rgba(0,0,0,0.1);
        }}
        .recommendation {{
            background-color: #f0f7ff;
            border-left: 5px solid #0066cc;
            padding: 15px;
            margin-bottom: 10px;
        }}
        .alert {{
            background-color: #fff8f8;
            border-left: 5px solid #cc0000;
            padding: 15px;
            margin-bottom: 10px;
        }}
        .info {{
            background-color: #f8f8ff;
            border-left: 5px solid #6666cc;
            padding: 15px;
            margin-bottom: 10px;
        }}
            </style>
        </head>
        <body>
    <div class="container">
        <h1>广东省空气质量监测报告</h1>
        
        <div class="header-info">
            <p>区域: {region if region != 'all' else '全省'}</p>
            <p>时间范围: {start_date} 至 {end_date}</p>
        </div>
        
        <div class="timestamp">
            <p>生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
        </div>
"""
        
        # 1. 空气质量概览部分
        if content_options.get('air_quality', True) and aqi_stats:
            html += """
        <div class="section">
            <h2>1. 空气质量概览</h2>
            
            <div class="stats-container">
"""
            
            # 添加AQI统计卡片
            for label, key, label_text in [
                ('数据记录数', 'count', '记录数'),
                ('平均AQI', 'avg', '平均AQI'),
                ('最高AQI', 'max', '最高AQI'),
                ('最低AQI', 'min', '最低AQI')
            ]:
                if key in aqi_stats:
                    html += f"""
                <div class="stats-card">
                    <div class="label">{label_text}</div>
                    <div class="value">{aqi_stats[key]}</div>
                </div>
"""
            
            html += """
            </div>
            
            <h3>空气质量数据表</h3>
            <table>
                <tr>
                    <th>日期</th>
                    <th>AQI</th>
                    <th>质量等级</th>
                    <th>城市</th>
                </tr>
"""
        
            # 添加AQI数据行（最多显示15行）
            for _, row in data['aqi_data'].head(15).iterrows():
                html += f"""
                <tr>
                    <td>{row.get('date')}</td>
                    <td>{row.get('aqi')}</td>
                    <td>{row.get('quality_level')}</td>
                    <td>{row.get('city')}</td>
                </tr>
"""
        
        html += """
            </table>
        </div>
"""
        
        # 2. 污染物分析部分
        if content_options.get('pollutants', True) and pollutants_stats:
            html += """
        <div class="section">
            <h2>2. 主要污染物分析</h2>
            
            <p>以下是监测期间各项污染物的平均浓度：</p>
            
            <div class="pollutant-grid">
"""
            
            # 添加污染物统计卡片
            for pollutant, value in pollutants_stats.items():
                html += f"""
                <div class="pollutant-card">
                    <div class="label">{pollutant}</div>
                    <div class="value">{value}</div>
                </div>
"""
            
            html += """
            </div>
        </div>
"""
        
        # 3. 趋势分析
        if content_options.get('trends', True):
            html += """
        <div class="section">
            <h2>3. 空气质量趋势</h2>
            
            <div class="info">
                <p>本报告HTML版暂不支持图表显示。未来版本将添加趋势图表功能。</p>
            </div>
        </div>
"""
        
        # 4. 预警信息
        if content_options.get('alerts', True) and not data['aqi_data'].empty:
            html += """
        <div class="section">
            <h2>4. 空气质量预警信息</h2>
"""
            
            # 检查是否有超标情况
            high_aqi = data['aqi_data'][data['aqi_data']['aqi'] > 100]
            if not high_aqi.empty:
                html += f"""
            <div class="alert">
                <p>发现 {len(high_aqi)} 天空气质量超过100，可能对敏感人群健康造成影响。</p>
            </div>
            
            <table>
                <tr>
                    <th>日期</th>
                    <th>AQI</th>
                    <th>质量等级</th>
                    <th>城市</th>
                </tr>
"""
                
                for _, row in high_aqi.head(10).iterrows():
                    html += f"""
                <tr>
                    <td>{row.get('date')}</td>
                    <td>{row.get('aqi')}</td>
                    <td>{row.get('quality_level')}</td>
                    <td>{row.get('city')}</td>
                </tr>
"""
                
                html += """
            </table>
"""
            else:
                html += """
            <div class="info">
                <p>监测期间内未发现明显空气质量超标情况。</p>
            </div>
"""
            
            html += """
        </div>
"""
        
        # 5. 政策建议
        if content_options.get('recommendations', False):
            html += """
        <div class="section">
            <h2>5. 改善空气质量的建议</h2>
"""
            
            recommendations = [
                "加强工业排放监管，严格执行排放标准",
                "推广清洁能源使用，减少化石燃料消耗",
                "加强城市绿化建设，增加空气净化能力",
                "提高公众环保意识，鼓励绿色出行方式"
            ]
            
            for rec in recommendations:
                html += f"""
            <div class="recommendation">
                <p>{rec}</p>
            </div>
"""
            
            html += """
        </div>
"""
        
        # 结束文档
        html += """
    </div>
        </body>
</html>
"""
        
        # 保存HTML文件
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html)
        
        return output_path
    except Exception as e:
        logger.error(f"HTML生成失败: {str(e)}\n{traceback.format_exc()}")
        return {'error': str(e)}

# 初始化服务
def init_reports_api():
    try:
        # 注册中文字体 
        font_path = os.path.join(project_root, 'data', 'fonts', 'simhei.ttf')
        if os.path.exists(font_path):
            pdfmetrics.registerFont(TTFont('SimHei', font_path))
        return True
    except Exception as e:
        logger.error(f"初始化失败: {str(e)}")
        return False

def get_report_type_name(report_type):
    """获取报告类型的中文名称
    
    Args:
        report_type (str): 报告类型代码
        
    Returns:
        str: 报告类型的中文名称
    """
    report_types = {
        'daily': '日报',
        'weekly': '周报',
        'monthly': '月报',
        'quarterly': '季度报告',
        'yearly': '年度报告',
        'custom': '自定义报告'
    }
    return report_types.get(report_type, '未知类型')

def get_region_name(region_code):
    """获取区域的中文名称
    
    Args:
        region_code (str): 区域代码
        
    Returns:
        str: 区域的中文名称
    """
    region_names = {
        'all': '全省',
        'guangzhou': '广州市',
        'shenzhen': '深圳市',
        'foshan': '佛山市',
        'dongguan': '东莞市',
        'zhongshan': '中山市',
        'zhuhai': '珠海市',
        'huizhou': '惠州市',
        'jiangmen': '江门市',
        'zhaoqing': '肇庆市',
        'shaoguan': '韶关市',
        'qingyuan': '清远市',
        'meizhou': '梅州市',
        'heyuan': '河源市',
        'shanwei': '汕尾市',
        'jieyang': '揭阳市',
        'maoming': '茂名市',
        'yangjiang': '阳江市',
        'zhanjiang': '湛江市',
        'chaozhou': '潮州市',
        'shantou': '汕头市',
        'yunfu': '云浮市',
        'pearl_delta': '珠三角',
        'east': '粤东地区',
        'west': '粤西地区',
        'north': '粤北地区'
    }
    return region_names.get(region_code, region_code)

@app.route('/api/reports/<report_id>', methods=['DELETE'])
def delete_report(report_id):
    """删除指定ID的报告
    
    Args:
        report_id (str): 报告ID
        
    Returns:
        JSON: 操作结果
    """
    try:
        if not report_id:
            logger.error("删除请求缺少报告ID")
            return handle_error("删除请求缺少报告ID", 400)
            
        # 记录删除请求
        logger.info(f"收到删除报告请求: {report_id}")
        
        # 从元数据中查找报告
        reports = load_report_metadata()
        if not reports:
            logger.warning("元数据为空，无法找到报告")
            return handle_error(f"找不到报告 {report_id}", 404)
            
        if not isinstance(reports, list):
            logger.error(f"元数据格式错误: {type(reports)}")
            return handle_error("报告元数据格式错误", 500)
            
        # 查找要删除的报告
        report_to_delete = None
        updated_reports = []
        
        for report in reports:
            if not isinstance(report, dict):
                continue
                
            # 检查ID是否匹配
            if report.get('id') == report_id:
                report_to_delete = report
            else:
                updated_reports.append(report)
                
        # 检查是否找到报告
        if not report_to_delete:
            logger.warning(f"找不到ID为 {report_id} 的报告")
            return handle_error(f"找不到报告 {report_id}", 404)
            
        # 获取报告文件路径
        file_path = report_to_delete.get('file_path')
        if file_path and os.path.exists(file_path):
            try:
                # 删除文件
                os.remove(file_path)
                logger.info(f"已删除报告文件: {file_path}")
            except Exception as e:
                logger.error(f"删除报告文件失败: {str(e)}")
                # 仍然继续删除元数据记录
        else:
            logger.warning(f"报告文件不存在或路径无效: {file_path}")
            
        # 更新元数据
        save_report_metadata(updated_reports)
        logger.info(f"已从元数据中移除报告 {report_id}")
        
        # 返回成功响应
        return jsonify({
            'success': True,
            'message': f"报告 {report_id} 已成功删除",
            'deleted_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        })
            
    except Exception as e:
        logger.error(f"删除报告失败: {str(e)}\n{traceback.format_exc()}")
        return handle_error(f"删除报告失败: {str(e)}")

# 启动服务
if __name__ == '__main__':
    init_reports_api()
    app.run(host='0.0.0.0', port=5003, debug=False)

# 添加报告预览API端点
@app.route('/api/reports/preview/<report_id>', methods=['GET'])
def preview_report(report_id):
    """生成报告预览内容
    
    Args:
        report_id: 报告ID
        
    Returns:
        HTML预览内容或错误信息
    """
    try:
        logger.info(f"收到报告预览请求: {report_id}")
        
        # 清除报告ID中可能包含的文件扩展名
        clean_id = report_id.split('.')[0] if '.' in report_id else report_id
        
        # 从元数据中查找报告记录
        reports = load_report_metadata()
        target_report = None
        
        for report in reports:
            if not isinstance(report, dict):
                continue
                
            if report.get('id') == clean_id:
                target_report = report
                break
        
        if not target_report:
            logger.error(f"没有找到报告ID: {clean_id}")
            return handle_error(f"报告预览失败: 没有找到报告ID {clean_id}", 404)
        
        # 获取报告文件路径
        file_path = target_report.get('file_path')
        
        if not file_path or not os.path.exists(file_path):
            logger.error(f"报告文件不存在: {file_path}")
            return handle_error("报告预览失败: 报告文件不存在", 404)
        
        # 根据文件格式生成预览内容
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.html':
            # 对于HTML格式，直接返回文件内容
            with open(file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
            return Response(html_content, mimetype='text/html')
        
        elif file_ext == '.pdf':
            # 对于PDF格式，生成一个HTML预览
            html_preview = f"""
            <html>
            <head>
                <title>报告预览 - {target_report.get('name', '未命名报告')}</title>
                <style>
                    body {{ font-family: Arial, sans-serif; margin: 0; padding: 20px; }}
                    .preview-info {{ background: #f5f5f5; padding: 15px; border-radius: 5px; margin-bottom: 20px; }}
                    .download-link {{ display: inline-block; padding: 10px 15px; background: #4CAF50; color: white; 
                                     text-decoration: none; border-radius: 4px; margin-top: 10px; }}
                    .pdf-container {{ width: 100%; height: 600px; border: 1px solid #ddd; }}
                </style>
            </head>
            <body>
                <div class="preview-info">
                    <h2>{target_report.get('name', '未命名报告')}</h2>
                    <p>格式: PDF | 生成时间: {target_report.get('created_at', '未知')}</p>
                    <p>因技术限制无法直接在浏览器中预览PDF内容，请下载查看完整报告。</p>
                    <a href="/api/reports/download/{target_report.get('id')}.pdf" class="download-link">下载报告</a>
                </div>
                <div>
                    <p>报告ID: {target_report.get('id')}</p>
                    <p>报告类型: {get_report_type_name(target_report.get('type', 'daily'))}</p>
                    <p>报告区域: {get_region_name(target_report.get('region', 'all'))}</p>
                    <p>时间范围: {target_report.get('start_date', '')} 至 {target_report.get('end_date', '')}</p>
                </div>
            </body>
            </html>
            """
            return Response(html_preview, mimetype='text/html')
        
        elif file_ext in ['.xlsx', '.xls']:
            # 对于Excel格式，生成一个HTML预览
            html_preview = f"""
            <html>
            <head>
                <title>报告预览 - {target_report.get('name', '未命名报告')}</title>
                <style>
                    body {{ font-family: Arial, sans-serif; margin: 0; padding: 20px; }}
                    .preview-info {{ background: #f5f5f5; padding: 15px; border-radius: 5px; margin-bottom: 20px; }}
                    .download-link {{ display: inline-block; padding: 10px 15px; background: #4CAF50; color: white; 
                                     text-decoration: none; border-radius: 4px; margin-top: 10px; }}
                </style>
            </head>
            <body>
                <div class="preview-info">
                    <h2>{target_report.get('name', '未命名报告')}</h2>
                    <p>格式: Excel | 生成时间: {target_report.get('created_at', '未知')}</p>
                    <p>Excel文件无法在浏览器中直接预览，请下载查看完整报告。</p>
                    <a href="/api/reports/download/{target_report.get('id')}.xlsx" class="download-link">下载报告</a>
                </div>
                <div>
                    <p>报告ID: {target_report.get('id')}</p>
                    <p>报告类型: {get_report_type_name(target_report.get('type', 'daily'))}</p>
                    <p>报告区域: {get_region_name(target_report.get('region', 'all'))}</p>
                    <p>时间范围: {target_report.get('start_date', '')} 至 {target_report.get('end_date', '')}</p>
                </div>
            </body>
            </html>
            """
            return Response(html_preview, mimetype='text/html')
        
        elif file_ext in ['.docx', '.doc']:
            # 对于Word格式，生成一个HTML预览
            html_preview = f"""
            <html>
            <head>
                <title>报告预览 - {target_report.get('name', '未命名报告')}</title>
                <style>
                    body {{ font-family: Arial, sans-serif; margin: 0; padding: 20px; }}
                    .preview-info {{ background: #f5f5f5; padding: 15px; border-radius: 5px; margin-bottom: 20px; }}
                    .download-link {{ display: inline-block; padding: 10px 15px; background: #4CAF50; color: white; 
                                     text-decoration: none; border-radius: 4px; margin-top: 10px; }}
                </style>
            </head>
            <body>
                <div class="preview-info">
                    <h2>{target_report.get('name', '未命名报告')}</h2>
                    <p>格式: Word | 生成时间: {target_report.get('created_at', '未知')}</p>
                    <p>Word文件无法在浏览器中直接预览，请下载查看完整报告。</p>
                    <a href="/api/reports/download/{target_report.get('id')}.docx" class="download-link">下载报告</a>
                </div>
                <div>
                    <p>报告ID: {target_report.get('id')}</p>
                    <p>报告类型: {get_report_type_name(target_report.get('type', 'daily'))}</p>
                    <p>报告区域: {get_region_name(target_report.get('region', 'all'))}</p>
                    <p>时间范围: {target_report.get('start_date', '')} 至 {target_report.get('end_date', '')}</p>
                </div>
            </body>
            </html>
            """
            return Response(html_preview, mimetype='text/html')
        
        else:
            # 对于其他格式，返回通用预览信息
            return handle_error(f"不支持预览格式: {file_ext}", 400)
    
    except Exception as e:
        logger.error(f"生成报告预览失败: {str(e)}\n{traceback.format_exc()}")
        return handle_error(f"生成报告预览失败: {str(e)}")